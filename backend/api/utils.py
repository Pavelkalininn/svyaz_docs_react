# -*- coding: UTF-8 -*-
from datetime import (
    date,
)
from docx import (
    Document,
)
from docx.shared import (
    Pt,
    RGBColor,
)


def value_or_zero(checked_obj, default_zero_value: str = '') -> str:
    if isinstance(checked_obj, date):
        return checked_obj.strftime("%d.%m.%Y")
    elif checked_obj is None:
        return default_zero_value
    return str(checked_obj)


def paragraph_checker(paragraph, change_line):
    if ('^' not in paragraph.text
            and paragraph.text.strip("<").strip(
                ">").strip("\'") in change_line.keys()):
        mark_in_paragraph = paragraph.text.strip(
            "<").strip(">").strip("\'")
        alignment = paragraph.alignment
        style = paragraph.style
        if paragraph.runs:
            font = paragraph.runs[0].font
            size = font.size
            font_name = font.name
            font_color = font.color.rgb
        else:
            size = Pt(10)
            font_name = 'Times New Roman'
            font_color = RGBColor(0x00, 0x00, 0x00)

        if change_line.get(mark_in_paragraph) is not None:
            if isinstance(
                    change_line.get(mark_in_paragraph), date
            ):
                temp_split_new_text = [
                    value_or_zero(change_line.get(mark_in_paragraph)), ]
            else:
                temp_split_new_text = change_line.get(
                    mark_in_paragraph).split("\n")
            for value in temp_split_new_text:
                new_par = paragraph.insert_paragraph_before(
                    '', style=style)
                new_par.alignment = alignment
                run = new_par.add_run(value_or_zero(value))
                run.font.size = size
                run.font.name = font_name
                run.font.color.rgb = font_color
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    else:
        if '^' in paragraph.text:
            mark_list = paragraph.text.split(
                '^')
            if len(mark_list) > 0:
                for mark in mark_list:
                    if mark in change_line.keys():
                        paragraph.text = (
                            paragraph.text.replace(
                                '^' + mark + '^',
                                value_or_zero(change_line.get(
                                    mark))
                            )
                        )


def table_and_paragraphs_checker(
        object_or_section: Document,
        change_line: dict
) -> None:
    for table in object_or_section.tables:
        for row in table.rows:
            for cell in row.cells:
                for table_paragraph in cell.paragraphs:
                    paragraph_checker(table_paragraph, change_line)
    for paragraph_in_obj in object_or_section.paragraphs:
        paragraph_checker(paragraph_in_obj, change_line)


def sections_checker(full_doc: Document, change_line: dict):
    for section in full_doc.sections:
        table_and_paragraphs_checker(section.header, change_line)
        table_and_paragraphs_checker(section.footer, change_line)
    table_and_paragraphs_checker(full_doc, change_line)


def obj_checker(checked_obj: Document, change_line: dict) -> None:
    sections_checker(checked_obj, change_line)
